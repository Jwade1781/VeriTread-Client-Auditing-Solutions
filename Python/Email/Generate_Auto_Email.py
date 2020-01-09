#######################################################################
#
# Purpose:
# The purpose of this program is to generate an email message that will
# be sent to the receiver. The generated email is stored in ..\Data\temp\Generated_Emails folder
# 
#
# TODO:
# [TBD]
#
# Extra Info:
# The message is stored as a .docx file, allowing for formating to take place.
# The message will be sent to the receiver by attaching it to the email
#
# Run:
# python ./Generate_Auto_Email.py dot_number
#
# EX:
# python ./Generate_Auto_Email.py 124
#######################################################################
def __main():
    dot_number = Parse_Command_Line()
    print("Generating an email message to be sent to DOT: " + str(dot_number))
    acceptance = Get_Acceptance(dot_number)
    message = Generate_Message(acceptance)
    document = Generate_Document(message)
    Save_Email_File(dot_number, document)
    return

#######################################################################
def Parse_Command_Line():
    # Parses the Command line to get the dot_number
    try:
        import sys
        dot_number = sys.argv[1]
        return dot_number

    except:
        print('No DOT number included in command line argument.. Exiting')
        exit(0)

#######################################################################
def Get_Acceptance(dot_number):
    import os
    path = os.path.dirname(os.path.realpath(__file__))
    import time
    time.sleep(5)

    file = open('..\Data\Saved_Predictions\\' + str(dot_number) + '.txt', 'r')
    date = file.readline()
    acceptance = file.readline()
    return acceptance

#######################################################################
def Generate_Message(acceptance):
    # Generate the message that will be placed inside the .docx file
    message = 'Hello, this is the Registration Services at VeriTread.\n'

    if (acceptance == 'Accepted'):
        message = message + 'We at VeriTread are pleased to announce that you were accepted to create an account on VeriTread!\n'

    elif (acceptance == 'Denied'):
        message = message + 'Thank you for you interest in registering with us.\n'
        message = message + 'Unfortuantely the automatic registration system at VeriTread has deemed that our services may not be a good match for your company based on the information we could find.\n'
        #message = message + 'If you believe that this is a mistake please contact us directly at (___) ___ - _____ or email@veritread.com\n'
        
    else:
        message = message + 'Based on the information that we were able to find, we will have to do a manual review to approve of your registration.\n'
        message = message + 'This can take some time to do and we may contact you directly for more information.\n'
        message = message + 'Thank you for your patience, and we will get back into contact with you as soon as possible.'

    return message

#######################################################################
def Generate_Document(message):
    # Genereates the .docx file based on the created message
    # This file will be added as an attachment to the email to be viewed
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    document = Document()
    document.add_picture('..\Data\Assets\\veritread_logo.png')
    document.add_heading('VeriTread Registration', 0)
    paragraph_1 = document.add_paragraph(message)

    contact_info_message = 'Have questions or concerns? Contact us at (___) ___ - ____ or VeriTread@veritread.com'
    paragraph_2 = document.add_paragraph(contact_info_message)
    return document

#######################################################################
def Save_Email_File(dot_number, document):        
    # Save the email .docx file into a temp folder
    path = '..\Data\\temp\Generated_Emails\\'
    file_name = path + str(dot_number) + '.docx'
    document.save(file_name)
    print("Stored Message into temp storage at: " + file_name)

__main()
